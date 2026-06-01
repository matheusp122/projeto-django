import argparse
import sys
from pathlib import Path

try:
    from PyPDF2 import PdfReader
except ImportError:
    print("Erro: a biblioteca PyPDF2 não está instalada. Instale com: pip install PyPDF2")
    sys.exit(1)

try:
    import docx
except ImportError:
    docx = None


def convert_pdf_to_text(pdf_path: Path, txt_path: Path) -> None:
    reader = PdfReader(str(pdf_path))

    text_lines = []
    for page_number, page in enumerate(reader.pages, start=1):
        try:
            page_text = page.extract_text() or ""
        except Exception as exc:
            raise RuntimeError(f"Falha ao extrair texto da página {page_number}: {exc}") from exc
        text_lines.append(page_text)

    txt_path.write_text("\n\n".join(text_lines), encoding="utf-8")


def convert_docx_to_text(docx_path: Path, txt_path: Path) -> None:
    if docx is None:
        raise RuntimeError(
            "Erro: a biblioteca python-docx não está instalada. Instale com: pip install python-docx"
        )

    document = docx.Document(str(docx_path))
    text_lines = []

    for paragraph in document.paragraphs:
        text_lines.append(paragraph.text)

    for table in document.tables:
        for row in table.rows:
            row_text = "\t".join(cell.text for cell in row.cells)
            text_lines.append(row_text)

    txt_path.write_text("\n\n".join(text_lines), encoding="utf-8")


def main() -> int:
    parser = argparse.ArgumentParser(
        description="Converte um arquivo PDF ou DOCX em um arquivo TXT de texto extraído."
    )
    parser.add_argument("input_file", help="Caminho do arquivo PDF ou DOCX de entrada")
    parser.add_argument(
        "txt_file",
        nargs="?",
        help="Caminho do arquivo TXT de saída (opcional). Se omitido, usa mesmo nome do arquivo de entrada.",
    )
    args = parser.parse_args()

    input_path = Path(args.input_file)
    if not input_path.exists():
        print(f"Erro: arquivo de entrada não encontrado: {input_path}")
        return 1

    extension = input_path.suffix.lower()
    if extension not in {".pdf", ".docx"}:
        print(f"Erro: o arquivo de entrada deve ser um PDF ou DOCX. Arquivo informado: {input_path}")
        return 1

    txt_path = Path(args.txt_file) if args.txt_file else input_path.with_suffix(".txt")
    try:
        if extension == ".pdf":
            convert_pdf_to_text(input_path, txt_path)
        else:
            convert_docx_to_text(input_path, txt_path)
    except Exception as exc:
        print(f"Erro ao converter {input_path.name} para TXT: {exc}")
        return 1

    print(f"Conversão concluída: {txt_path}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())